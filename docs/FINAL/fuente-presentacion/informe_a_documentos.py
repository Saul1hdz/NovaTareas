# -*- coding: utf-8 -*-
"""Convierte informe-final.md a DOCX editable y PDF, desde un mismo análisis.

Un solo parser alimenta los dos generadores: así los dos documentos dicen lo
mismo, en el mismo orden, con la misma jerarquía. Si divergen, es por cómo
cada formato dibuja, no por el contenido.
"""
import io
import os
import re
import sys

ORIGEN = sys.argv[1]
SALIDA = os.path.dirname(ORIGEN)
BASE = 'informe-final'

# ── Análisis del Markdown ────────────────────────────────────────────────────
# Bloques: {'tipo': 'h1'|'h2'|'h3'|'p'|'lista'|'numerada'|'tabla'|'codigo'|'regla'}

def parsear(texto):
    lineas = texto.split('\n')
    bloques = []
    i = 0
    parrafo = []

    def cerrar_parrafo():
        if parrafo:
            bloques.append({'tipo': 'p', 'texto': ' '.join(parrafo).strip()})
            parrafo.clear()

    while i < len(lineas):
        linea = lineas[i]
        despojada = linea.strip()

        if despojada.startswith('```'):
            cerrar_parrafo()
            i += 1
            codigo = []
            while i < len(lineas) and not lineas[i].strip().startswith('```'):
                codigo.append(lineas[i])
                i += 1
            bloques.append({'tipo': 'codigo', 'lineas': codigo})
            i += 1
            continue

        if re.match(r'^---+$', despojada):
            cerrar_parrafo()
            bloques.append({'tipo': 'regla'})
            i += 1
            continue

        encabezado = re.match(r'^(#{1,4})\s+(.*)$', despojada)
        if encabezado:
            cerrar_parrafo()
            nivel = len(encabezado.group(1))
            bloques.append({'tipo': f'h{min(nivel, 3)}', 'texto': encabezado.group(2).strip()})
            i += 1
            continue

        # Tabla: una línea con | seguida de la fila separadora |---|
        if despojada.startswith('|') and i + 1 < len(lineas) and re.match(r'^\|[\s:|-]+\|$', lineas[i + 1].strip()):
            cerrar_parrafo()
            def celdas(fila):
                return [c.strip() for c in fila.strip().strip('|').split('|')]
            cabecera = celdas(lineas[i])
            i += 2
            filas = []
            while i < len(lineas) and lineas[i].strip().startswith('|'):
                filas.append(celdas(lineas[i]))
                i += 1
            bloques.append({'tipo': 'tabla', 'cabecera': cabecera, 'filas': filas})
            continue

        vinieta = re.match(r'^[-*]\s+(.*)$', despojada)
        numerada = re.match(r'^(\d+)\.\s+(.*)$', despojada)
        if vinieta or numerada:
            cerrar_parrafo()
            tipo = 'lista' if vinieta else 'numerada'
            items = []
            while i < len(lineas):
                actual = lineas[i].strip()
                v = re.match(r'^[-*]\s+(.*)$', actual)
                n = re.match(r'^(\d+)\.\s+(.*)$', actual)
                coincide = v if tipo == 'lista' else n
                if coincide:
                    items.append(coincide.group(1) if tipo == 'lista' else coincide.group(2))
                    i += 1
                elif actual and lineas[i].startswith(('  ', '\t')) and items:
                    items[-1] += ' ' + actual        # continuación del punto anterior
                    i += 1
                else:
                    break
            bloques.append({'tipo': tipo, 'items': items})
            continue

        if despojada.startswith('>'):
            cerrar_parrafo()
            cita = []
            while i < len(lineas) and lineas[i].strip().startswith('>'):
                cita.append(lineas[i].strip().lstrip('>').strip())
                i += 1
            bloques.append({'tipo': 'cita', 'texto': ' '.join(cita).strip()})
            continue

        if not despojada:
            cerrar_parrafo()
        else:
            parrafo.append(despojada)
        i += 1

    cerrar_parrafo()
    return bloques


# ── Texto en línea: negrita, código y enlaces ────────────────────────────────
PATRON_NEGRITA = re.compile(r'(\*\*.+?\*\*)')
PATRON_INLINE = re.compile(r'(`[^`]+`|\[[^\]]+\]\([^)]+\)|<https?://[^>]+>)')

def _sin_negrita(texto):
    """Código y enlaces dentro de un fragmento ya sin marcas de negrita."""
    salida = []
    for parte in PATRON_INLINE.split(texto):
        if not parte:
            continue
        if parte.startswith('`') and parte.endswith('`'):
            salida.append((parte[1:-1], {'codigo'}))
        elif parte.startswith('<http'):
            # Enlace automatico: se muestra la direccion, sin los angulos.
            salida.append((parte[1:-1], {'enlace'}))
        elif parte.startswith('['):
            # Puede no ser un enlace: los emoji sustituidos ([OK], [Parcial])
            # también empiezan por corchete y no deben tratarse como tales.
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', parte)
            if not m:
                salida.append((parte, set()))
                continue
            etiqueta, destino = m.group(1), m.group(2)
            # Una etiqueta como `archivo.md` conserva su aire de código.
            estilo = {'enlace'} if destino.startswith('http') else {'codigo'}
            salida.append((etiqueta.strip('`'), estilo))
        else:
            salida.append((parte, set()))
    return salida


def trozos(texto):
    """Devuelve [(texto, estilos)] con estilos ⊆ {negrita, codigo, enlace}.

    El análisis es en dos pasadas porque los estilos se combinan: un
    **`comando`** en negrita es a la vez negrita y código, y resolverlo con una
    sola expresión regular dejaba los backticks a la vista.
    """
    salida = []
    for parte in PATRON_NEGRITA.split(texto):
        if not parte:
            continue
        if parte.startswith('**') and parte.endswith('**'):
            for contenido, estilos in _sin_negrita(parte[2:-2]):
                salida.append((contenido, estilos | {'negrita'}))
        else:
            salida.extend(_sin_negrita(parte))
    return salida or [(texto, set())]


def plano(texto):
    return ''.join(t for t, _ in trozos(texto))


# Los emoji no existen en las fuentes del PDF: se sustituyen por su significado.
EQUIVALENCIAS = {
    '✅': '[OK]', '🟡': '[Parcial]', '❌': '[Pendiente]', '⚠️': '[Aviso]', '⚠': '[Aviso]',
    '👍': 'positiva', '👎': 'negativa', '⬜': '[Pendiente]', '✦': '',
}

def sin_emoji(texto):
    for emoji, reemplazo in EQUIVALENCIAS.items():
        texto = texto.replace(emoji, reemplazo)
    return re.sub(r'\s{2,}', ' ', texto).replace('positiva/negativa', 'positiva o negativa').strip()


# ── DOCX ─────────────────────────────────────────────────────────────────────
def generar_docx(bloques, destino):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor, Inches

    doc = Document()
    for seccion in doc.sections:
        seccion.top_margin = seccion.bottom_margin = Inches(0.9)
        seccion.left_margin = seccion.right_margin = Inches(0.9)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    def escribir(parrafo, texto):
        for contenido, estilos in trozos(texto):
            run = parrafo.add_run(contenido)
            if 'negrita' in estilos:
                run.bold = True
            if 'codigo' in estilos:
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
            if 'enlace' in estilos:
                run.font.color.rgb = RGBColor(0x1A, 0x4F, 0xB0)
                run.underline = True

    for bloque in bloques:
        tipo = bloque['tipo']

        if tipo == 'h1':
            p = doc.add_heading(level=0)
            escribir(p, bloque['texto'])
        elif tipo in ('h2', 'h3'):
            p = doc.add_heading(level=2 if tipo == 'h2' else 3)
            escribir(p, bloque['texto'])
        elif tipo == 'p':
            escribir(doc.add_paragraph(), bloque['texto'])
        elif tipo == 'cita':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            escribir(p, bloque['texto'])
            for run in p.runs:
                run.italic = True
        elif tipo in ('lista', 'numerada'):
            estilo = 'List Bullet' if tipo == 'lista' else 'List Number'
            for item in bloque['items']:
                escribir(doc.add_paragraph(style=estilo), item)
        elif tipo == 'codigo':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(10)
            for indice, linea in enumerate(bloque['lineas']):
                run = p.add_run(linea)
                run.font.name = 'Consolas'
                run.font.size = Pt(8.5)
                if indice < len(bloque['lineas']) - 1:
                    run.add_break()
        elif tipo == 'tabla':
            columnas = len(bloque['cabecera'])
            con_cabecera = any(c.strip() for c in bloque['cabecera'])
            tabla = doc.add_table(rows=1, cols=columnas)
            tabla.style = 'Light Grid Accent 1' if con_cabecera else 'Light List'
            if con_cabecera:
                for celda, titulo in zip(tabla.rows[0].cells, bloque['cabecera']):
                    celda.text = ''
                    escribir(celda.paragraphs[0], titulo)
                    for run in celda.paragraphs[0].runs:
                        run.bold = True
            else:
                tabla._tbl.remove(tabla.rows[0]._tr)
            for fila in bloque['filas']:
                celdas = tabla.add_row().cells
                for celda, valor in zip(celdas, fila):
                    celda.text = ''
                    escribir(celda.paragraphs[0], valor)
            doc.add_paragraph()
        elif tipo == 'regla':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(destino)


# ── PDF ──────────────────────────────────────────────────────────────────────
def generar_pdf(bloques, destino):
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (BaseDocTemplate, Frame, PageTemplate, Paragraph,
                                    Preformatted, Spacer, Table, TableStyle, KeepTogether)
    from xml.sax.saxutils import escape

    FUENTES = 'C:/Windows/Fonts/'
    pdfmetrics.registerFont(TTFont('Texto', FUENTES + 'segoeui.ttf'))
    pdfmetrics.registerFont(TTFont('Texto-Neg', FUENTES + 'segoeuib.ttf'))
    pdfmetrics.registerFont(TTFont('Texto-Cur', FUENTES + 'segoeuii.ttf'))
    pdfmetrics.registerFont(TTFont('Mono', FUENTES + 'consola.ttf'))
    pdfmetrics.registerFont(TTFont('Mono-Neg', FUENTES + 'consolab.ttf'))
    pdfmetrics.registerFontFamily('Texto', normal='Texto', bold='Texto-Neg', italic='Texto-Cur')

    TINTA = colors.HexColor('#1a1a1a')
    ACENTO = colors.HexColor('#4b3fd4')
    SUAVE = colors.HexColor('#5c5c66')

    estilos = {
        'h1': ParagraphStyle('h1', fontName='Texto-Neg', fontSize=20, leading=25,
                             textColor=ACENTO, spaceAfter=10, spaceBefore=0),
        'h2': ParagraphStyle('h2', fontName='Texto-Neg', fontSize=13.5, leading=18,
                             textColor=ACENTO, spaceBefore=16, spaceAfter=7),
        'h3': ParagraphStyle('h3', fontName='Texto-Neg', fontSize=11, leading=14.5,
                             textColor=TINTA, spaceBefore=11, spaceAfter=5),
        'p': ParagraphStyle('p', fontName='Texto', fontSize=9.5, leading=14,
                            textColor=TINTA, spaceAfter=7, alignment=TA_LEFT),
        'cita': ParagraphStyle('cita', fontName='Texto-Cur', fontSize=9, leading=13.5,
                               textColor=SUAVE, leftIndent=14, spaceAfter=8),
        'item': ParagraphStyle('item', fontName='Texto', fontSize=9.5, leading=13.5,
                               textColor=TINTA, leftIndent=16, bulletIndent=4, spaceAfter=3),
        'celda': ParagraphStyle('celda', fontName='Texto', fontSize=8.3, leading=11,
                                textColor=TINTA),
        'celda-cab': ParagraphStyle('celda-cab', fontName='Texto-Neg', fontSize=8.3,
                                    leading=11, textColor=colors.white),
        'codigo': ParagraphStyle('codigo', fontName='Mono', fontSize=7.6, leading=10.2,
                                 textColor=TINTA, leftIndent=8),
    }

    def marcado(texto, mono='Mono'):
        """Markdown en línea → etiquetas de reportlab."""
        salida = []
        for contenido, estilos in trozos(sin_emoji(texto)):
            seguro = escape(contenido)
            if 'codigo' in estilos:
                seguro = f'<font face="{mono}" size="8.2">{seguro}</font>'
            elif 'enlace' in estilos:
                seguro = f'<font color="#1a4fb0">{seguro}</font>'
            if 'negrita' in estilos:
                seguro = f'<b>{seguro}</b>'
            salida.append(seguro)
        return ''.join(salida)

    ANCHO = LETTER[0] - 2 * 2.2 * cm
    historia = []

    for bloque in bloques:
        tipo = bloque['tipo']
        if tipo in ('h1', 'h2', 'h3'):
            historia.append(Paragraph(marcado(bloque['texto']), estilos[tipo]))
        elif tipo == 'p':
            historia.append(Paragraph(marcado(bloque['texto']), estilos['p']))
        elif tipo == 'cita':
            historia.append(Paragraph(marcado(bloque['texto']), estilos['cita']))
        elif tipo in ('lista', 'numerada'):
            for indice, item in enumerate(bloque['items'], 1):
                vineta = '•' if tipo == 'lista' else f'{indice}.'
                historia.append(Paragraph(marcado(item), estilos['item'], bulletText=vineta))
            historia.append(Spacer(1, 5))
        elif tipo == 'codigo':
            texto = '\n'.join(bloque['lineas'])
            caja = Preformatted(texto, estilos['codigo'])
            fondo = Table([[caja]], colWidths=[ANCHO])
            fondo.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f4f4f8')),
                ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor('#dcdce4')),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 7),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
            ]))
            historia.extend([fondo, Spacer(1, 9)])
        elif tipo == 'tabla':
            columnas = len(bloque['cabecera'])
            con_cabecera = any(c.strip() for c in bloque['cabecera'])
            datos = []
            if con_cabecera:
                datos.append([Paragraph(marcado(c), estilos['celda-cab']) for c in bloque['cabecera']])
            for fila in bloque['filas']:
                fila = (fila + [''] * columnas)[:columnas]
                datos.append([Paragraph(marcado(c), estilos['celda']) for c in fila])
            # La primera columna suele llevar la etiqueta y las demás el dato.
            if columnas == 2:
                anchos = [ANCHO * 0.34, ANCHO * 0.66]
            else:
                primera = ANCHO * (0.34 if columnas <= 3 else 0.28)
                resto = (ANCHO - primera) / (columnas - 1)
                anchos = [primera] + [resto] * (columnas - 1)
            tabla = Table(datos, colWidths=anchos, repeatRows=1 if con_cabecera else 0)
            primera = 1 if con_cabecera else 0
            tabla.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), ACENTO if con_cabecera else colors.white),
                ('ROWBACKGROUNDS', (0, primera), (-1, -1), [colors.white, colors.HexColor('#f6f6fa')]),
                ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#d5d5de')),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 5),
                ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ]))
            historia.extend([tabla, Spacer(1, 10)])
        elif tipo == 'regla':
            historia.append(Spacer(1, 4))

    def pie(canvas, doc):
        canvas.saveState()
        canvas.setFont('Texto', 7.5)
        canvas.setFillColor(SUAVE)
        canvas.drawString(2.2 * cm, 1.3 * cm, 'NovaTareas Pro 1.0.0 — Informe final')
        canvas.drawRightString(LETTER[0] - 2.2 * cm, 1.3 * cm, f'Página {doc.page}')
        canvas.restoreState()

    documento = BaseDocTemplate(destino, pagesize=LETTER,
                                leftMargin=2.2 * cm, rightMargin=2.2 * cm,
                                topMargin=1.9 * cm, bottomMargin=2.0 * cm,
                                title='Informe final — NovaTareas Pro',
                                author='Equipo 3 — Universidad Gerardo Barrios')
    marco = Frame(documento.leftMargin, documento.bottomMargin,
                  documento.width, documento.height, id='cuerpo')
    documento.addPageTemplates([PageTemplate(id='normal', frames=[marco], onPage=pie)])
    documento.build(historia)


texto = io.open(ORIGEN, encoding='utf-8').read()
bloques = parsear(texto)
resumen = {}
for b in bloques:
    resumen[b['tipo']] = resumen.get(b['tipo'], 0) + 1
print('bloques:', resumen)

generar_docx(bloques, os.path.join(SALIDA, BASE + '.docx'))
print('DOCX listo')
generar_pdf(bloques, os.path.join(SALIDA, BASE + '.pdf'))
print('PDF listo')
